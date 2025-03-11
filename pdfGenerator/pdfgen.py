import subprocess
import sys
try:
    import docx
except ImportError:
    print('Không tìm thấy python-docx! Đang tự động cài...')
    subprocess.check_call([sys.executable, "-m", "pip", "install", 'python-docx'])
try:
    import docx2pdf
except ImportError:
    print('Không tìm thấy docx2pdf! Đang tự động cài...')
    subprocess.check_call([sys.executable, "-m", "pip", "install", 'docx2pdf'])

import datetime
from datetime import datetime,date,timedelta
import time
import random
#get user info
print('Tool tạo PDF gậy v1.0\nĐược viết vào ngày 11/3/2025 bởi @realEtheriaa\nLưu ý: Thông tin của bạn có thể fake, nhưng thông tin tác giả phải chính xác để đảm bảo gậy đáp trên 95%! (Có thể tra Google/Wikipedia)')
user = input('>>> Nhập tên của bạn (đầy đủ có dấu): ')
fm = int(input('>>> Nhập 1 = nam, 2 = nữ: '))
dob = input('>>> Nhập ngày tháng năm sinh (dd/mm/YYYY): ')
phone = input(">>> Số điện thoại của bạn: ")
proc_dob = datetime.strptime(dob, "%d/%m/%Y").date()
male=0
female=1
if proc_dob.year <= 2000:
    male=0
    female=1
elif proc_dob.year <= 2010:
    male=2
    female=3
elif proc_dob.year <= 2020:
    male=4
    female=5
elif proc_dob <= 2030:
    male=6
    female=7
elif proc_dob <= 2040:
    male=8
    female=9

province = [
  { "ma_tinh": "01", "ten_tinh": "Hà Nội" },
  { "ma_tinh": "02", "ten_tinh": "Hà Giang" },
  { "ma_tinh": "04", "ten_tinh": "Cao Bằng" },
  { "ma_tinh": "06", "ten_tinh": "Bắc Kạn" },
  { "ma_tinh": "08", "ten_tinh": "Tuyên Quang" },
  { "ma_tinh": "10", "ten_tinh": "Lào Cai" },
  { "ma_tinh": "11", "ten_tinh": "Điện Biên" },
  { "ma_tinh": "12", "ten_tinh": "Lai Châu" },
  { "ma_tinh": "14", "ten_tinh": "Sơn La" },
  { "ma_tinh": "15", "ten_tinh": "Yên Bái" },
  { "ma_tinh": "17", "ten_tinh": "Hòa Bình" },
  { "ma_tinh": "19", "ten_tinh": "Thái Nguyên" },
  { "ma_tinh": "20", "ten_tinh": "Lạng Sơn" },
  { "ma_tinh": "22", "ten_tinh": "Quảng Ninh" },
  { "ma_tinh": "24", "ten_tinh": "Bắc Giang" },
  { "ma_tinh": "25", "ten_tinh": "Phú Thọ" },
  { "ma_tinh": "26", "ten_tinh": "Vĩnh Phúc" },
  { "ma_tinh": "27", "ten_tinh": "Bắc Ninh" },
  { "ma_tinh": "30", "ten_tinh": "Hải Dương" },
  { "ma_tinh": "31", "ten_tinh": "Hải Phòng" },
  { "ma_tinh": "33", "ten_tinh": "Hưng Yên" },
  { "ma_tinh": "34", "ten_tinh": "Thái Bình" },
  { "ma_tinh": "35", "ten_tinh": "Hà Nam" },
  { "ma_tinh": "36", "ten_tinh": "Nam Định" },
  { "ma_tinh": "37", "ten_tinh": "Ninh Bình" },
  { "ma_tinh": "38", "ten_tinh": "Thanh Hóa" },
  { "ma_tinh": "40", "ten_tinh": "Nghệ An" },
  { "ma_tinh": "42", "ten_tinh": "Hà Tĩnh" },
  { "ma_tinh": "44", "ten_tinh": "Quảng Bình" },
  { "ma_tinh": "45", "ten_tinh": "Quảng Trị" },
  { "ma_tinh": "46", "ten_tinh": "Thừa Thiên Huế" },
  { "ma_tinh": "48", "ten_tinh": "Đà Nẵng" },
  { "ma_tinh": "49", "ten_tinh": "Quảng Nam" },
  { "ma_tinh": "51", "ten_tinh": "Quảng Ngãi" },
  { "ma_tinh": "52", "ten_tinh": "Bình Định" },
  { "ma_tinh": "54", "ten_tinh": "Phú Yên" },
  { "ma_tinh": "56", "ten_tinh": "Khánh Hòa" },
  { "ma_tinh": "58", "ten_tinh": "Ninh Thuận" },
  { "ma_tinh": "60", "ten_tinh": "Bình Thuận" },
  { "ma_tinh": "62", "ten_tinh": "Kon Tum" },
  { "ma_tinh": "64", "ten_tinh": "Gia Lai" },
  { "ma_tinh": "66", "ten_tinh": "Đắk Lắk" },
  { "ma_tinh": "67", "ten_tinh": "Đắk Nông" },
  { "ma_tinh": "68", "ten_tinh": "Lâm Đồng" },
  { "ma_tinh": "70", "ten_tinh": "Bình Phước" },
  { "ma_tinh": "72", "ten_tinh": "Tây Ninh" },
  { "ma_tinh": "74", "ten_tinh": "Bình Dương" },
  { "ma_tinh": "75", "ten_tinh": "Đồng Nai" },
  { "ma_tinh": "77", "ten_tinh": "Bà Rịa - Vũng Tàu" },
  { "ma_tinh": "79", "ten_tinh": "TP. Hồ Chí Minh" },
  { "ma_tinh": "80", "ten_tinh": "Long An" },
  { "ma_tinh": "82", "ten_tinh": "Tiền Giang" },
  { "ma_tinh": "83", "ten_tinh": "Bến Tre" },
  { "ma_tinh": "84", "ten_tinh": "Trà Vinh" },
  { "ma_tinh": "86", "ten_tinh": "Vĩnh Long" },
  { "ma_tinh": "87", "ten_tinh": "Đồng Tháp" },
  { "ma_tinh": "89", "ten_tinh": "An Giang" },
  { "ma_tinh": "91", "ten_tinh": "Kiên Giang" },
  { "ma_tinh": "92", "ten_tinh": "Cần Thơ" },
  { "ma_tinh": "93", "ten_tinh": "Hậu Giang" },
  { "ma_tinh": "94", "ten_tinh": "Sóc Trăng" },
  { "ma_tinh": "95", "ten_tinh": "Bạc Liêu" },
  { "ma_tinh": "96", "ten_tinh": "Cà Mau" }
]
print(">>> Dưới đây là danh sách mã tỉnh: ")
for i in province:
    print(i["ma_tinh"], " ------", i["ten_tinh"])
liveprov = ""
boole = True
while boole:
    live = input('>>> Nhập mã tỉnh nơi bạn sống: ')
    for i in province:
        if str(live) == str(i["ma_tinh"]):
            liveprov = live
            boole = False
            break
    if boole:
        print("Mã tỉnh không hợp lệ!")
y = abs(proc_dob.year) % 100
randomIDnum = str(0) + str(liveprov) + str(male if fm % 2 != 0 else female) + str(y if y > 10 else str(0) + str(y)) + str(random.randint(0, 99999)).rjust(6, "0")
print(">>> ! <<<    Đây là mã CCCD fake của bạn: " + randomIDnum)
mst = liveprov + str(random.randint(0,9999999)).rjust(8,"0")
addr = input(">>> Nhập địa chỉ mà bạn nghĩ ra (phải cùng tỉnh với tỉnh bạn đã chọn): ")
def str_time_prop(start, end, time_format, prop):
    stime = time.mktime(time.strptime(start, time_format))
    etime = time.mktime(time.strptime(end, time_format))
    ptime = stime + prop * (etime - stime)
    return time.strftime(time_format, time.localtime(ptime))
def random_date(start, end, prop):
    return str_time_prop(start, end, '%d/%m/%Y', prop)
toda = date.today()
past = toda + timedelta(days=-4*365)
mailg = input(">>> Nhập email gậy của bạn: ")

user2 = input('>>> Nhập tên của nghệ sĩ (đầy đủ có dấu): ')
user3 = input('>>> Nhập nghệ danh của nghệ sĩ (VD: MONO, ...): ')
fm2 = int(input('>>> Nhập 1 = nam, 2 = nữ: '))
dob2 = input('>>> Nhập ngày tháng năm sinh (dd/mm/YYYY): ')
proc_dob2 = datetime.strptime(dob, "%d/%m/%Y").date()
male=0
female=1
if proc_dob2.year <= 2000:
    male=0
    female=1
elif proc_dob2.year <= 2010:
    male=2
    female=3
elif proc_dob2.year <= 2020:
    male=4
    female=5
elif proc_dob2.year <= 2030:
    male=6
    female=7
elif proc_dob2.year <= 2040:
    male=8
    female=9
liveprov2 = ""
boole2 = True
while boole2:
    live = input('>>> Nhập mã tỉnh nơi nghệ sĩ sống: ')
    for i in province:
        if str(live) == str(i["ma_tinh"]):
            liveprov2 = live
            boole2 = False
            break
    if boole2:
        print("Mã tỉnh không hợp lệ!")
y = abs(proc_dob2.year) % 100
randomIDnum2 = str(0) + str(liveprov2) + str(male if fm2 % 2 != 0 else female) + str(y if y > 10 else str(0) + str(y)) + str(random.randint(0, 99999)).rjust(6, "0")
print(">>> ! <<<    Đây là mã CCCD fake của nghệ sĩ: " + randomIDnum2)
mst2 = liveprov2 + str(random.randint(0,9999999)).rjust(8,"0")
print(">>> ! <<<    Đây là mã số thuế fake của nghệ sĩ: " + mst2)

prod_name = input(">>> Nhập tên tác phẩm:")
rel_date = input(">>> Nhập ngày phát hành tác phẩm (dd/mm/yyyy): ")
alias2 = input(">>> Nhập đầy đủ nghệ danh của các nghệ sĩ tham gia (VD: 'AMEE & MCK', 'Lương Quý Tuấn & Hữu Công', ...)")
print('\nĐợi chút nhé, tool đang tạo file...')
proc_rdate = datetime.strptime(rel_date, "%d/%m/%Y")
signed_date = None
if proc_rdate >= datetime.strptime("30/09/2024", "%d/%m/%Y"):
    signed_date = random_date(rel_date, date.today(), random.random())
else:
    signed_date = random_date("30/09/2024", date.today().strftime('%d/%m/%Y'), random.random())

#process
#Skipping CRC checking
import zipfile
import tempfile
# Override the _update_crc method with a no-op to ignore CRC errors
def no_crc_update(self, data):
    pass
zipfile.ZipExtFile._update_crc = no_crc_update
#DOCX Processing
from docx import Document
import urllib.request
temp = tempfile.NamedTemporaryFile()
url = "https://fuchsia.viiic.net/sample.docx"
# Your credentials
username = 'ngdepchai00112'
password = 'watafacisbluddoing'
# Create a password manager and add your credentials
password_mgr = urllib.request.HTTPPasswordMgrWithDefaultRealm()
password_mgr.add_password(None, url, username, password)

# Create an HTTP Basic Auth handler using the password manager
auth_handler = urllib.request.HTTPBasicAuthHandler(password_mgr)

# Build an opener that will use our auth handler
opener = urllib.request.build_opener(auth_handler)

# Use the opener to fetch the URL; the credentials will be sent in the Authorization header
with opener.open(url) as response:
    content = response.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(content)
        temp_file_path = tmp_file.name
doc = Document(temp_file_path)
replace_word = {
    'date1': datetime.strptime(signed_date, "%d/%m/%Y").strftime("ngày %d tháng %m năm %Y"),
    'date2': signed_date,
    'user1': user,
    'user2': user2,
    'alias1': user3,
    'numb1': phone,
    'cccd1': randomIDnum,
    'cccd2': randomIDnum2,
    'dob2': dob2,
    'mst1': mst,
    'mst2': mst2,
    'capngay1': random_date(past.strftime('%d/%m/%Y'), toda.strftime('%d/%m/%Y'), random.random()),
    'capngay2': random_date(past.strftime('%d/%m/%Y'), toda.strftime('%d/%m/%Y'), random.random()),
    'addr1': addr,
    'email': mailg,
    'prod_name': prod_name,
    'alias2': alias2
}

def replace_in_paragraph(paragraph, replacements):
    # Merge all runs into a single string to detect split placeholders
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Check if any placeholder exists
    modified = False
    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
            modified = True
    
    if modified:
        # Clear existing runs and add new text while retaining formatting
        for run in paragraph.runs:
            run.text = ""
        new_run = paragraph.add_run(full_text)
        # Copy font style from the first run (optional)
        if paragraph.runs:
            new_run.font.name = paragraph.runs[0].font.name
            new_run.font.bold = paragraph.runs[0].font.bold

# Process paragraphs
for p in doc.paragraphs:
    replace_in_paragraph(p, replace_word)

# Process paragraphs
for p in doc.paragraphs:
    replace_in_paragraph(p, replace_word)

# Process tables and hyperlinks
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # Process cell paragraphs
            for p in cell.paragraphs:
                replace_in_paragraph(p, replace_word)

fname = prod_name + ' - ' + alias2
doc.save(fname+'.docx')

from docx2pdf import convert
convert(fname+'.docx',fname+'.pdf')